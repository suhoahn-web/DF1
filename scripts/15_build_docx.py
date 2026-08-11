"""Assemble the submission document (ESWA Your-Paper-Your-Way single file):
manuscript body + figures with captions + generated tables.

Input:  라이팅/01_MANUSCRIPT.md, 라이팅/04_TABLES.md, results/figures/*.png
Output: 라이팅/CGRC_submission.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
MS = ROOT / "라이팅" / "01_MANUSCRIPT.md"
OUT = ROOT / "라이팅" / (sys.argv[1] if len(sys.argv) > 1 else "CGRC_submission.docx")

FIGS = [
    ("fig1_timeline", "Figure 1. Temporal design: expanding base retraining with strictly "
     "ordered TCN-training, gate-training, development-fold, and holdout periods."),
    ("fig2_gate_behavior", "Figure 2. Learned gate behavior on both bases: (a) mean gate by "
     "context; (b) mean gate by recent base-error decile."),
    ("fig3_regime_delta", "Figure 3. Pooled WAPE change of the gated system versus its base, "
     "by regime and base (development folds)."),
    ("fig4_case_study", "Figure 4. Case study (TFT base): actual sales, base forecast, gated "
     "forecast, and gate values around promotion episodes (shaded)."),
]


def add_markdown(doc: Document, text: str) -> None:
    skip_sections = {"Appendices (assembled from results/)"}
    skipping = False
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line.lstrip("#").strip()
            skipping = title in skip_sections
            if skipping:
                continue
            if level == 1:
                doc.add_heading(title, level=0)
            else:
                doc.add_heading(title, level=min(level - 1, 3))
        elif skipping or line.startswith("**Authors") or line.startswith("Authors:") \
                or line.startswith("Target:") or line == "---":
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line:
            doc.add_paragraph(line)


def add_tables(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Tables", level=1)
    tables_md = (ROOT / "라이팅" / "04_TABLES.md").read_text(encoding="utf-8")
    blocks = re.split(r"^## ", tables_md, flags=re.M)[1:]
    for block in blocks:
        lines = [ln for ln in block.splitlines() if ln.strip()]
        caption, rows = lines[0], [ln for ln in lines[1:] if ln.startswith("|")]
        doc.add_paragraph(caption).runs[0].bold = True
        cells = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
        cells = [r for r in cells if not all(set(c) <= {"-", ":", ""} for c in r)]
        if not cells:
            continue
        t = doc.add_table(rows=len(cells), cols=len(cells[0]))
        t.style = "Light Grid Accent 1"
        for i, row in enumerate(cells):
            for j, val in enumerate(row[:len(t.columns)]):
                t.rows[i].cells[j].text = val
        doc.add_paragraph("")


def add_figures(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Figures", level=1)
    for name, caption in FIGS:
        p = ROOT / "results" / "figures" / f"{name}.png"
        if p.exists():
            doc.add_picture(str(p), width=Inches(6.2))
            doc.add_paragraph(caption)
            doc.add_paragraph("")


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    add_markdown(doc, MS.read_text(encoding="utf-8"))
    add_tables(doc)
    add_figures(doc)
    doc.save(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    main()
